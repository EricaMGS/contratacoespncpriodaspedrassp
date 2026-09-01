"""
Painel de Inteligência e Apoio ao Controle Interno - PNCP
Prefeitura Municipal de Rio das Pedras/SP

Tecnologias:
- Streamlit
- Pandas
- Requests
- Scikit-Learn (opcional)
- OpenPyXL
- python-docx (necessário somente para exportação Word)

Objetivo:
Consulta, consolidação, segmentação, análise preliminar de risco,
monitoramento de vigência, busca ativa de termos aditivos e exportação.

IMPORTANTE:
Os alertas produzidos pelo sistema são indicadores automatizados.
Toda conclusão de auditoria deve ser validada por servidor responsável.
"""

from __future__ import annotations

import io
import logging
import os
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

import pandas as pd
import requests
import streamlit as st

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Pt
    DOCX_OK = True
except ImportError:
    Document = None
    WD_ALIGN_PARAGRAPH = None
    WD_TABLE_ALIGNMENT = None
    Pt = None
    DOCX_OK = False

from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry


# ============================================================
# CONFIGURAÇÃO DE LOG
# ============================================================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
)

LOGGER = logging.getLogger(__name__)


# ============================================================
# DEPENDÊNCIAS OPCIONAIS
# ============================================================

try:
    from sklearn.feature_extraction.text import TfidfVectorizer

    SKLEARN_OK = True
except ImportError:
    TfidfVectorizer = None
    SKLEARN_OK = False

    LOGGER.warning(
        "scikit-learn não encontrado. "
        "A funcionalidade de similaridade semântica será desativada."
    )


# ============================================================
# CONFIGURAÇÕES GLOBAIS
# ============================================================

CNPJ = "44826840000183"
IBGE = "3544004"

MUNICIPIO = "Rio das Pedras/SP"
PREFEITURA = "Prefeitura Municipal de Rio das Pedras/SP"

BASE_CONSULTA = "https://pncp.gov.br/api/consulta/v1"
# API de serviços do PNCP: contratos, termos, histórico e documentos.
BASE_SERVICOS = "https://pncp.gov.br/api/pncp/v1"

MAX_WORKERS = max(
    1,
    min(
        16,
        int(os.getenv("PNCP_MAX_WORKERS", "6")),
    ),
)

TIMEOUT_CONNECT = max(
    3,
    int(os.getenv("PNCP_TIMEOUT_CONNECT", "10")),
)

TIMEOUT_READ = max(
    10,
    int(os.getenv("PNCP_TIMEOUT_READ", "45")),
)

TIMEOUT = (TIMEOUT_CONNECT, TIMEOUT_READ)

MAX_TENTATIVAS = max(
    1,
    int(os.getenv("PNCP_MAX_RETRIES", "4")),
)

PAGE_CONTRATOS = 100
PAGE_ATAS = 100
PAGE_EDITAIS = 50

MAX_PAGINAS = max(
    1,
    int(os.getenv("PNCP_MAX_PAGES", "15")),
)

CACHE_TTL = 900

TIPOS = [
    "Contratos",
    "Atas de Registro de Preços",
    "Editais e Avisos de Contratações",
]

MODALIDADES_PNCP = tuple(range(1, 16))

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 "
        "Chrome/131.0 Safari/537.36 "
        "MonitoramentoControleInterno/6.2"
    ),
    "Accept": "application/json",
    "Accept-Encoding": "gzip, deflate",
    "Connection": "keep-alive",
}


# ============================================================
# ESTILOS
# ============================================================

def aplicar_estilo() -> None:
    """Aplica pequenos ajustes visuais sem alterar a identidade do app."""

    st.markdown(
        """
        <style>
        .stMetric {
            border-radius: 10px;
        }

        div[data-testid="stMetricValue"] {
            font-size: 1.55rem;
        }

        .risk-high {
            color: #b91c1c;
            font-weight: 700;
        }

        .risk-medium {
            color: #b45309;
            font-weight: 700;
        }

        .risk-low {
            color: #15803d;
            font-weight: 700;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


# ============================================================
# HTTP
# ============================================================

@st.cache_resource(show_spinner=False)
def sessao_http() -> requests.Session:
    """Cria uma sessão HTTP reutilizável."""

    session = requests.Session()
    session.headers.update(HEADERS)

    retry = Retry(
        total=MAX_TENTATIVAS,
        connect=MAX_TENTATIVAS,
        read=MAX_TENTATIVAS,
        status=MAX_TENTATIVAS,
        backoff_factor=1.5,
        status_forcelist=(408, 425, 429, 500, 502, 503, 504),
        allowed_methods=frozenset({"GET", "HEAD", "OPTIONS"}),
        respect_retry_after_header=True,
        raise_on_status=False,
    )

    adapter = HTTPAdapter(
        max_retries=retry,
        pool_connections=MAX_WORKERS,
        pool_maxsize=MAX_WORKERS,
    )

    session.mount("http://", adapter)
    session.mount("https://", adapter)

    return session


def detalhe_http(response: requests.Response) -> str:
    """Extrai uma mensagem útil da resposta HTTP."""

    try:
        payload = response.json()

        if isinstance(payload, dict):
            for key in (
                "message",
                "error",
                "detail",
                "mensagem",
                "erro",
            ):
                value = payload.get(key)

                if value is not None:
                    return str(value)[:500]

    except (ValueError, TypeError):
        pass

    text = (response.text or "").strip()

    return text[:500] if text else "Resposta sem detalhes."


def get_json(
    url: str,
    params: Optional[Dict[str, Any]] = None,
) -> Any:
    """Executa GET e retorna JSON."""

    session = sessao_http()

    try:
        response = session.get(
            url,
            params=params,
            timeout=TIMEOUT,
        )

    except requests.exceptions.Timeout as exc:
        raise RuntimeError(
            "Tempo limite excedido ao consultar o PNCP."
        ) from exc

    except requests.exceptions.ConnectionError as exc:
        raise RuntimeError(
            "Não foi possível estabelecer conexão com o PNCP."
        ) from exc

    except requests.exceptions.RequestException as exc:
        raise RuntimeError(
            f"Erro de comunicação com o PNCP: {exc}"
        ) from exc

    if response.status_code in (204, 404):
        return []

    if response.status_code == 200:
        try:
            return response.json()
        except ValueError as exc:
            raise RuntimeError(
                "O PNCP respondeu com conteúdo que não pôde ser interpretado como JSON."
            ) from exc

    detalhe = detalhe_http(response)

    if response.status_code == 429:
        raise RuntimeError(
            "O PNCP informou limite de requisições (HTTP 429). "
            "Aguarde alguns instantes e tente novamente."
        )

    if response.status_code in (500, 502, 503, 504):
        raise RuntimeError(
            f"O PNCP está temporariamente indisponível "
            f"(HTTP {response.status_code}). Detalhe: {detalhe}"
        )

    if response.status_code in (400, 422):
        raise RuntimeError(
            f"Parâmetros rejeitados pelo PNCP "
            f"(HTTP {response.status_code}): {detalhe}"
        )

    if response.status_code == 403:
        raise RuntimeError(
            "O PNCP recusou o acesso à consulta (HTTP 403)."
        )

    raise RuntimeError(
        f"PNCP respondeu HTTP {response.status_code}: {detalhe}"
    )


# ============================================================
# UTILITÁRIOS
# ============================================================

INVALID_TEXTS = {
    "",
    "none",
    "nan",
    "null",
    "n/d",
    "nat",
}


def valor_vazio(valor: Any) -> bool:
    """Determina se um valor pode ser considerado vazio."""

    if valor is None:
        return True

    if isinstance(valor, (dict, list, tuple, set)):
        return len(valor) == 0

    try:
        resultado = pd.isna(valor)

        if isinstance(resultado, bool):
            return resultado

    except (TypeError, ValueError):
        pass

    return False


def texto(
    valor: Any,
    padrao: str = "N/D",
) -> str:
    """Converte um valor em texto seguro."""

    if valor_vazio(valor):
        return padrao

    if isinstance(valor, (dict, list, tuple, set)):
        return str(valor)

    resultado = str(valor).strip()

    if resultado.lower() in INVALID_TEXTS:
        return padrao

    return resultado


def numero(valor: Any) -> Optional[float]:
    """Converte números brasileiros e internacionais para float."""

    if valor_vazio(valor) or isinstance(valor, bool):
        return None

    if isinstance(valor, (int, float)):
        try:
            resultado = float(valor)

            if pd.isna(resultado):
                return None

            return resultado

        except (TypeError, ValueError):
            return None

    s = str(valor).strip()

    if not s:
        return None

    s = (
        s.replace("R$", "")
        .replace("r$", "")
        .replace(" ", "")
        .replace("\xa0", "")
    )

    if "," in s and "." in s:
        ultima_virgula = s.rfind(",")
        ultimo_ponto = s.rfind(".")

        if ultima_virgula > ultimo_ponto:
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", "")

    elif "," in s:
        partes = s.split(",")
        if len(partes) == 2 and len(partes[1]) <= 2:
            s = s.replace(",", ".")
        else:
            s = s.replace(",", "")

    try:
        resultado = float(s)
        if pd.isna(resultado):
            return None
        return resultado
    except (TypeError, ValueError):
        return None


def moeda(valor: Any) -> str:
    """Formata um número como moeda brasileira."""

    n = numero(valor)

    if n is None:
        return "N/D"

    return (
        f"R$ {n:,.2f}"
        .replace(",", "X")
        .replace(".", ",")
        .replace("X", ".")
    )


def data_br(valor: Any) -> str:
    """Converte data para dd/mm/aaaa."""

    if valor_vazio(valor):
        return "N/D"

    data = pd.to_datetime(
        valor,
        errors="coerce",
    )

    if pd.isna(data):
        return "N/D"

    return data.strftime("%d/%m/%Y")


def data_timestamp(valor: Any) -> Optional[pd.Timestamp]:
    """Converte valor para pd.Timestamp seguro."""
    if valor_vazio(valor):
        return None
    ts = pd.to_datetime(valor, errors="coerce")
    return ts.normalize() if pd.notna(ts) else None


def cnpj_limpo(valor: Any) -> str:
    """Mantém somente números de um CNPJ."""

    if valor_vazio(valor):
        return ""

    return re.sub(r"\D", "", str(valor))


def slug(valor: Any) -> str:
    """Gera nome seguro para arquivos."""

    s = re.sub(
        r"[^A-Za-z0-9_-]+",
        "_",
        texto(valor, "registro"),
    )

    s = s.strip("_")

    return s[:100] or "registro"


def primeiro(
    row_dict: Dict[str, Any],
    campos: Sequence[str],
    padrao: Any = "N/D",
) -> Any:
    """Retorna o primeiro campo válido."""

    for campo in campos:
        if campo not in row_dict:
            continue

        valor = row_dict.get(campo)

        if valor_vazio(valor):
            continue

        if isinstance(valor, str):
            valor_limpo = valor.strip()

            if valor_limpo.lower() in INVALID_TEXTS:
                continue

            return valor_limpo

        return valor

    return padrao


def fuzzy_match(
    row_dict: Dict[str, Any],
    palavras: Sequence[str],
    padrao: Any = None
) -> Any:
    """Busca aproximada pelo nome do campo como fallback."""

    for chave, valor in row_dict.items():

        if valor_vazio(valor):
            continue

        chave_lower = str(chave).lower()

        if not any(
            palavra.lower() in chave_lower
            for palavra in palavras
        ):
            continue

        if isinstance(valor, (dict, list)):
            continue

        valor_str = str(valor).strip()

        if valor_str.lower() in INVALID_TEXTS:
            continue

        return valor

    return padrao


# ============================================================
# EXTRAÇÃO DOS REGISTROS COM BUSCA ATIVA DE TERMOS ADITIVOS
# ============================================================

def buscar_termos_contrato(
    cnpj_orgao: str,
    ano: Any,
    sequencial: Any,
) -> Dict[str, Any]:
    """Consulta os termos vinculados a um contrato e preserva o diagnóstico real."""

    resultado = {
        "status": "SEM_IDENTIFICADORES",
        "termos": [],
        "http_status": None,
        "mensagem": "",
        "url": "",
    }

    if valor_vazio(ano) or valor_vazio(sequencial):
        resultado["mensagem"] = (
            "O contrato não possui anoContrato e/ou sequencialContrato "
            "disponíveis para consulta dos termos."
        )
        return resultado

    cnpj = cnpj_limpo(cnpj_orgao)
    if not cnpj:
        resultado["mensagem"] = "CNPJ do órgão não disponível para consulta."
        return resultado

    try:
        ano_int = int(ano)
        seq_int = int(sequencial)
    except (TypeError, ValueError):
        resultado["mensagem"] = (
            f"Identificadores inválidos: ano={ano!r}, sequencial={sequencial!r}."
        )
        return resultado

    url = (
        f"{BASE_SERVICOS}/orgaos/{cnpj}"
        f"/contratos/{ano_int}/{seq_int}/termos"
    )
    resultado["url"] = url

    session = sessao_http()

    try:
        response = session.get(url, timeout=TIMEOUT)
    except requests.exceptions.Timeout:
        resultado["status"] = "ERRO_CONSULTA"
        resultado["mensagem"] = "Tempo limite excedido ao consultar os termos no PNCP."
        LOGGER.warning("Timeout na consulta de termos %s/%s", ano_int, seq_int)
        return resultado
    except requests.exceptions.ConnectionError:
        resultado["status"] = "ERRO_CONSULTA"
        resultado["mensagem"] = "Não foi possível estabelecer conexão com o PNCP."
        LOGGER.warning("Erro de conexão na consulta de termos %s/%s", ano_int, seq_int)
        return resultado
    except requests.exceptions.RequestException as exc:
        resultado["status"] = "ERRO_CONSULTA"
        resultado["mensagem"] = f"Erro de comunicação com o PNCP: {exc}"
        LOGGER.warning("Erro de requisição na consulta de termos %s/%s: %s", ano_int, seq_int, exc)
        return resultado

    resultado["http_status"] = response.status_code

    if response.status_code == 200:
        try:
            payload = response.json()
        except ValueError:
            resultado["status"] = "RESPOSTA_INVALIDA"
            resultado["mensagem"] = (
                "O PNCP respondeu HTTP 200, mas o conteúdo não pôde ser "
                "interpretado como JSON."
            )
            return resultado

        termos = registros_api(payload)

        if not termos:
            resultado["status"] = "SEM_REGISTROS"
            resultado["mensagem"] = (
                "Consulta realizada com sucesso. O PNCP não retornou "
                "termos vinculados a este contrato."
            )
            return resultado

        resultado["status"] = "REGISTROS_ENCONTRADOS"
        resultado["termos"] = termos
        resultado["mensagem"] = f"{len(termos)} termo(s) localizado(s) no PNCP."
        return resultado

    if response.status_code == 204:
        resultado["status"] = "SEM_REGISTROS"
        resultado["mensagem"] = (
            "O PNCP respondeu HTTP 204 (sem conteúdo). Nenhum termo foi retornado."
        )
        return resultado

    if response.status_code == 404:
        resultado["status"] = "HTTP_404"
        resultado["mensagem"] = (
            "O PNCP respondeu HTTP 404 para o recurso de termos. "
            "Isso não deve ser interpretado automaticamente como ausência de aditivos."
        )
        LOGGER.warning("HTTP 404 na consulta de termos %s/%s", ano_int, seq_int)
        return resultado

    if response.status_code == 429:
        resultado["status"] = "ERRO_CONSULTA"
        resultado["mensagem"] = (
            "O PNCP informou limite de requisições (HTTP 429). "
            "A consulta dos termos não pôde ser concluída."
        )
        return resultado

    if response.status_code in (500, 502, 503, 504):
        resultado["status"] = "ERRO_CONSULTA"
        resultado["mensagem"] = (
            f"O PNCP apresentou indisponibilidade temporária "
            f"(HTTP {response.status_code})."
        )
        return resultado

    resultado["status"] = "ERRO_CONSULTA"
    resultado["mensagem"] = (
        f"O PNCP respondeu HTTP {response.status_code}: {detalhe_http(response)}"
    )
    return resultado


def dados_registro(
    row: Any,
    tipo: str,
) -> Dict[str, Any]:
    """Extrai os campos com checagem do sub-recurso de termos aditivos."""

    if hasattr(row, "to_dict"):
        r_dict = row.to_dict()
    else:
        r_dict = dict(row)

    controle = primeiro(
        r_dict,
        [
            "numeroControlePNCP",
            "numeroControlePNCPAta",
            "numeroControlePNCPCompra",
            "idContratoPNCP",
        ],
    )

    ano_contrato = primeiro(r_dict, ["anoContrato", "ano", "compra.anoCompra"], padrao=None)
    seq_contrato = primeiro(r_dict, ["sequencialContrato", "sequencial", "numeroSequencial"], padrao=None)

    tipo_instrumento = primeiro(
        r_dict,
        [
            "tipoContrato.nome",
            "tipoContratoNome",
            "tipoTermoContrato.nome",
            "tipoTermoContratoNome",
            "tipoDocumento.nome",
        ],
        padrao="Contrato",
    )

    tipo_contrato_id = primeiro(
        r_dict,
        [
            "tipoContrato.id",
            "tipoContratoId",
            "tipoTermoContrato.id",
            "tipoTermoContratoId",
        ],
        padrao=None,
    )

    dt_fim_vig = None
    val_inicial = None
    val_global = None

    if tipo == "Contratos":
        num = primeiro(r_dict, ["numeroContratoEmpenho", "numeroContrato", "numero"])
        proc = primeiro(r_dict, ["processo", "numeroProcesso", "compra.processo"])
        obj = primeiro(r_dict, ["objetoContrato", "objetoCompra", "compra.objetoCompra", "objeto", "descricaoObjeto"])
        forn = primeiro(r_dict, ["nomeRazaoSocialFornecedor", "fornecedor.nomeRazaoSocial", "razaoSocialFornecedor", "nomeFornecedor"])
        cnpj = primeiro(r_dict, ["niFornecedor", "fornecedor.niFornecedor", "cnpjFornecedor", "cnpj"])
        
        val_global = primeiro(r_dict, ["valorGlobal", "valorTotal", "valorContrato", "compra.valorTotalHomologado"])
        val_inicial = primeiro(r_dict, ["valorInicial", "valorTotalEstimado", "valorContrato"])
        
        dt_ = primeiro(r_dict, ["dataAssinatura", "dataCelebracao", "dataPublicacaoPncp"])
        dt_fim_vig = primeiro(r_dict, ["dataVigenciaFim", "dataFimVigencia", "dataTerminoVigencia", "dataEncerramento"])
        sit = primeiro(r_dict, ["situacao", "status"])

    elif tipo == "Atas de Registro de Preços":
        num = primeiro(r_dict, ["numeroAtaRegistroPreco", "numeroAta", "numero"])
        proc = primeiro(r_dict, ["processo", "numeroProcesso", "processoAdministrativo", "compra.processo"])
        obj = primeiro(r_dict, ["objetoAta", "objetoAtaRegistroPreco", "objetoCompra", "compra.objetoCompra", "objeto", "descricaoObjeto"])
        forn = primeiro(r_dict, ["nomeRazaoSocialFornecedor", "fornecedor.nomeRazaoSocial", "razaoSocialFornecedor", "nomeFornecedor"])
        cnpj = primeiro(r_dict, ["niFornecedor", "fornecedor.niFornecedor", "cnpjFornecedor", "ni"])
        
        val_global = primeiro(r_dict, ["valorTotalAta", "valorTotal", "valorGlobal", "valorAta", "compra.valorTotalHomologado"])
        val_inicial = primeiro(r_dict, ["valorInicial", "valorTotalEstimado"])
        
        dt_ = primeiro(r_dict, ["dataAssinatura", "dataPublicacaoPncp", "dataCelebracao"])
        dt_fim_vig = primeiro(r_dict, ["dataVigenciaFim", "dataFimVigencia", "dataValidadeFim"])
        sit = primeiro(r_dict, ["situacao", "status"])

    else:
        num = primeiro(r_dict, ["numeroCompra", "compra.numeroCompra", "numeroEdital", "numero"])
        proc = primeiro(r_dict, ["processo", "compra.processo", "numeroProcesso"])
        obj = primeiro(r_dict, ["objetoCompra", "compra.objetoCompra", "objeto", "descricaoObjeto"])
        forn = primeiro(r_dict, ["nomeRazaoSocialFornecedor", "fornecedor.nomeRazaoSocial", "razaoSocialFornecedor", "nomeFornecedor"])
        cnpj = primeiro(r_dict, ["niFornecedor", "fornecedor.niFornecedor", "cnpjFornecedor", "cnpjOrgao"])
        
        val_global = primeiro(r_dict, ["valorTotalHomologado", "compra.valorTotalHomologado", "valorTotal", "valorGlobal"])
        val_inicial = primeiro(r_dict, ["valorTotalEstimado", "compra.valorTotalEstimado"])
        
        dt_ = primeiro(r_dict, ["dataPublicacao", "dataPublicacaoPncp", "dataInclusao"])
        dt_fim_vig = primeiro(r_dict, ["dataVigenciaFim", "dataFimVigencia"])
        sit = primeiro(r_dict, ["situacaoCompra", "compra.situacaoCompraNome", "situacao", "status"])

    mod = primeiro(r_dict, ["modalidadeNome", "modalidadeContratacaoNome", "compra.modalidadeNome", "modalidade"])

    if texto(obj, "") == "":
        obj = fuzzy_match(r_dict, ["objeto", "descricao"])

    if texto(forn, "") == "":
        forn = fuzzy_match(r_dict, ["razaosocial", "fornec", "fornecedor.nome"])

    if texto(val_global, "") == "":
        val_global = fuzzy_match(r_dict, ["valortotal", "valorglobal", "valorata", "valor"])

    if dt_fim_vig == "N/D" or dt_fim_vig is None:
        dt_fim_vig = fuzzy_match(r_dict, ["vigenciafim", "fimvigencia", "datatermino", "datavalidadefim"])

    if texto(mod, "") == "":
        mod = fuzzy_match(r_dict, ["modalidade", "tipoata"])

    if texto(cnpj, "") == "":
        cnpj = fuzzy_match(r_dict, ["cnpj", "ni"])

    v_num = numero(val_global)
    v_ini_num = numero(val_inicial)

    # 1. Checagem inicial no próprio payload
    eh_termo_aditivo = False
    if str(tipo_contrato_id) == "2" or "aditivo" in str(tipo_instrumento).lower():
        eh_termo_aditivo = True
    elif any(palavra in str(obj).lower() for palavra in ("termo aditivo", "aditamento")):
        eh_termo_aditivo = True

    # 2. Busca ativa nos sub-recursos de Termos do PNCP para contratos
    termos_vinculados = []
    diagnostico_aditivo = {
        "status": "NAO_CONSULTADO",
        "termos": [],
        "http_status": None,
        "mensagem": "Consulta de termos não executada para este registro.",
        "url": "",
    }

    if tipo == "Contratos":
        diagnostico_aditivo = buscar_termos_contrato(
            CNPJ,
            ano_contrato,
            seq_contrato,
        )
        termos_vinculados = diagnostico_aditivo.get("termos", [])

        if diagnostico_aditivo.get("status") == "REGISTROS_ENCONTRADOS":
            eh_termo_aditivo = True

    # 3. Cálculo de variação percentual de aditivos
    perc_aditivo = None
    if v_num is not None and v_ini_num is not None and v_ini_num > 0 and v_num > v_ini_num:
        perc_aditivo = ((v_num - v_ini_num) / v_ini_num) * 100.0

    ts_fim_vig = data_timestamp(dt_fim_vig)
    dias_para_vencer = None
    if ts_fim_vig is not None:
        hoje = pd.Timestamp.today().normalize()
        dias_para_vencer = int((ts_fim_vig - hoje).days)

    return {
        "controle": texto(controle),
        "numero": texto(num),
        "processo": texto(proc),
        "ano_contrato": ano_contrato,
        "sequencial_contrato": seq_contrato,
        "tipo_instrumento": texto(tipo_instrumento),
        "eh_aditivo": eh_termo_aditivo,
        "qtd_termos_aditivos": len(termos_vinculados),
        "termos_detalhes": termos_vinculados,
        "diagnostico_aditivo": diagnostico_aditivo,
        "objeto": texto(obj),
        "fornecedor": texto(forn),
        "cnpj_fornecedor": cnpj_limpo(cnpj),
        "valor_num": v_num,
        "valor_inicial_num": v_ini_num,
        "perc_aditivo": perc_aditivo,
        "valor": moeda(val_global),
        "data": data_br(dt_),
        "data_fim_vigencia": data_br(dt_fim_vig),
        "dias_para_vencer": dias_para_vencer,
        "situacao": texto(sit),
        "modalidade": texto(mod),
    }


# ============================================================
# PAGINAÇÃO
# ============================================================

def registros_api(data: Any) -> List[Dict[str, Any]]:
    """Extrai a lista de registros de diferentes formatos possíveis."""

    if isinstance(data, list):
        return [item for item in data if isinstance(item, dict)]

    if not isinstance(data, dict):
        return []

    for chave in ("data", "items", "content", "dados", "registros"):
        valor = data.get(chave)
        if isinstance(valor, list):
            return [item for item in valor if isinstance(item, dict)]

    return []


def inteiro_seguro(valor: Any) -> Optional[int]:
    """Converte um valor para inteiro sem lançar exceção."""
    if valor_vazio(valor):
        return None
    try:
        return int(valor)
    except (TypeError, ValueError):
        return None


def paginacao(data: Any) -> Dict[str, Optional[int]]:
    """Extrai informações de paginação."""
    if not isinstance(data, dict):
        return {"totalPaginas": None, "totalRegistros": None, "numeroPagina": None}

    return {
        "totalPaginas": inteiro_seguro(data.get("totalPaginas")),
        "totalRegistros": inteiro_seguro(data.get("totalRegistros")),
        "numeroPagina": inteiro_seguro(data.get("numeroPagina")),
    }


def calcular_total_paginas(
    info: Dict[str, Optional[int]],
    tamanho_pagina: int,
) -> Optional[int]:
    """Determina o total de páginas."""
    total_paginas = info.get("totalPaginas")
    if total_paginas is not None:
        return max(1, total_paginas)

    total_registros = info.get("totalRegistros")
    if total_registros is None:
        return None

    tamanho = max(1, tamanho_pagina)
    return max(1, (total_registros + tamanho - 1) // tamanho)


@st.cache_data(ttl=CACHE_TTL, show_spinner=False)
def consultar_cache(
    url: str,
    params_tuple: Tuple[Tuple[str, str], ...],
    max_paginas: int,
) -> Tuple[List[Dict[str, Any]], int, Optional[int], List[str]]:
    """Consulta paginada com paralelismo."""
    base = dict(params_tuple)
    base["pagina"] = 1

    primeira = get_json(url, base)
    registros_primeira = registros_api(primeira)
    info = paginacao(primeira)

    tamanho_pagina = inteiro_seguro(base.get("tamanhoPagina")) or 50
    total = calcular_total_paginas(info, tamanho_pagina)

    if not registros_primeira:
        return ([], 1, total, [])

    limite = min(max(1, max_paginas), total if total is not None else max_paginas)
    todos = list(registros_primeira)

    if limite <= 1:
        return (todos, 1, total, [])

    resultados: Dict[int, List[Dict[str, Any]]] = {}
    erros: List[str] = []
    workers = min(MAX_WORKERS, max(1, limite - 1))

    def buscar_pagina(pagina: int) -> Tuple[int, List[Dict[str, Any]]]:
        params = dict(base)
        params["pagina"] = pagina
        resposta = get_json(url, params)
        return pagina, registros_api(resposta)

    with ThreadPoolExecutor(max_workers=workers) as executor:
        futuros = {executor.submit(buscar_pagina, p): p for p in range(2, limite + 1)}

        for futuro in as_completed(futuros):
            pagina = futuros[futuro]
            try:
                pagina_resultado, registros = futuro.result()
                resultados[pagina_resultado] = registros
            except Exception as exc:
                mensagem = f"Página {pagina}: {exc}"
                LOGGER.error("Erro na consulta: %s", mensagem)
                erros.append(mensagem)

    for pagina in range(2, limite + 1):
        todos.extend(resultados.get(pagina, []))

    return (todos, limite, total, erros)


def consultar(
    url: str,
    params: Dict[str, Any],
    max_paginas: int,
) -> Tuple[List[Dict[str, Any]], int, Optional[int], List[str]]:
    """Wrapper da consulta com tupla serializada para o cache."""
    serial = tuple(sorted((str(k), str(v)) for k, v in params.items()))
    return consultar_cache(url, serial, max_paginas)


# ============================================================
# DATAFRAME
# ============================================================

def normalizar_pncp(registros: List[Dict[str, Any]]) -> pd.DataFrame:
    """Normaliza JSON PNCP para DataFrame."""
    if not registros:
        return pd.DataFrame()

    df = pd.json_normalize(registros, sep=".")
    if df.empty:
        return df

    for coluna in df.columns:
        df[coluna] = df[coluna].map(lambda v: str(v) if isinstance(v, (dict, list)) else v)

    return df


def deduplicar(df: pd.DataFrame) -> pd.DataFrame:
    """Deduplica mantendo sequenciais diferentes de contratos/aditivos."""
    if df.empty:
        return df

    resultado = df.copy()

    colunas_compostas = []
    for c in ["numeroControlePNCP", "numeroControlePNCPAta", "numeroControlePNCPCompra", "idContratoPNCP"]:
        if c in resultado.columns:
            colunas_compostas.append(c)

    for c in ["sequencialContrato", "sequencialTermoContrato", "numeroTermoContrato", "tipoContrato.id"]:
        if c in resultado.columns:
            colunas_compostas.append(c)

    if colunas_compostas:
        resultado = resultado.drop_duplicates(subset=colunas_compostas, keep="first")
    else:
        resultado = resultado.drop_duplicates()

    return resultado.reset_index(drop=True)


def filtrar_cnpj(df: pd.DataFrame) -> pd.DataFrame:
    """Filtra o órgão pelo CNPJ."""
    if df.empty:
        return df

    alvo = cnpj_limpo(CNPJ)
    colunas = (
        "cnpjOrgao",
        "cnpj",
        "cnpjCompra",
        "orgaoEntidade.cnpj",
        "orgao.cnpj",
    )

    mascaras: List[pd.Series] = []

    for coluna in colunas:
        if coluna not in df.columns:
            continue
        serie = df[coluna].map(cnpj_limpo)
        mascara = serie == alvo
        if mascara.any():
            mascaras.append(mascara)

    if not mascaras:
        return df

    mascara_final = mascaras[0].copy()
    for mascara in mascaras[1:]:
        mascara_final |= mascara

    return df.loc[mascara_final].reset_index(drop=True)


# ============================================================
# FILTRO DE DATA
# ============================================================

COLUNAS_DATA = (
    "dataPublicacaoPncp",
    "dataPublicacao",
    "dataInclusao",
    "dataAssinatura",
    "dataCelebracao",
)


def filtrar_periodo(
    df: pd.DataFrame,
    inicio: pd.Timestamp,
    fim: pd.Timestamp,
) -> pd.DataFrame:
    """Aplica filtro temporal."""
    if df.empty:
        return df

    colunas_disponiveis = [c for c in COLUNAS_DATA if c in df.columns]
    if not colunas_disponiveis:
        return df

    mascaras = []
    for coluna in colunas_disponiveis:
        datas = pd.to_datetime(df[coluna], errors="coerce").dt.normalize()
        mascaras.append(datas.notna() & (datas >= inicio) & (datas <= fim))

    mascara_final = mascaras[0].copy()
    for mascara in mascaras[1:]:
        mascara_final |= mascara

    return df.loc[mascara_final].reset_index(drop=True)


# ============================================================
# MOTOR DE RISCO
# ============================================================

def limites_iqr(
    serie: pd.Series,
) -> Tuple[Optional[float], Optional[float], Optional[float]]:
    """Calcula limites IQR."""
    valores = pd.to_numeric(serie, errors="coerce").dropna()
    if len(valores) < 5:
        return None, None, None

    q1 = valores.quantile(0.25)
    q3 = valores.quantile(0.75)
    iqr = q3 - q1
    mediana = valores.median()

    if iqr <= 0:
        return float(q1), None, float(mediana)

    return float(q1 - 1.5 * iqr), float(q3 + 1.5 * iqr), float(mediana)


def calcular_risco(
    registro: Dict[str, Any],
    contexto: pd.DataFrame,
    tipo: str,
) -> Dict[str, Any]:
    """Calcula um score heurístico de risco enriquecido com vigência e aditivos."""
    pontos = 0
    motivos: List[str] = []
    testes: List[str] = []

    valor = registro.get("valor_num")
    modalidade = texto(registro.get("modalidade"), "").lower()
    objeto = texto(registro.get("objeto"), "").lower()

    # 1. Completude
    campos_obrigatorios = (("objeto", "objeto"), ("controle", "controle PNCP"), ("data", "data"))
    faltantes = [
        label for campo, label in campos_obrigatorios
        if texto(registro.get(campo), "") in {"", "N/D"}
    ]
    if faltantes:
        pontos += min(12, 4 * len(faltantes))
        motivos.append("Campos incompletos: " + ", ".join(faltantes))
        testes.append("Atenção na completude cadastral")

    # 2. Modalidade / Exceção
    if "dispensa" in modalidade or "inexig" in modalidade:
        pontos += 20
        motivos.append("Contratação por exceção")
        testes.append("Revisar fundamento legal")
    elif "emerg" in objeto or "emerg" in modalidade:
        pontos += 25
        motivos.append("Indício emergencial")
        testes.append("Revisar caracterização e prazo emergencial")

    # 3. Valor
    if valor is not None:
        if valor >= 500_000:
            pontos += 10
            motivos.append("Valor elevado (> R$ 500 mil)")
        elif valor >= 150_000:
            pontos += 5
            motivos.append("Valor relevante (> R$ 150 mil)")

    # 4. Outlier
    serie_valores = contexto["valor_num"] if "valor_num" in contexto.columns else pd.Series(dtype="float64")
    _, limite_superior, _ = limites_iqr(serie_valores)
    outlier = False
    if valor is not None and limite_superior is not None and valor > limite_superior:
        outlier = True
        pontos += 25
        motivos.append("Valor anômalo acima do limite IQR")
        testes.append("Avaliar formação de preços e pesquisa de mercado")

    # 5. Monitoramento de Vigência
    dias_venc = registro.get("dias_para_vencer")
    if dias_venc is not None:
        if 0 <= dias_venc <= 30:
            pontos += 15
            motivos.append(f"⏳ Vencimento crítico ({dias_venc} dias restantes)")
            testes.append("Verificar abertura tempestiva de nova licitação ou prorrogação")
        elif 30 < dias_venc <= 90:
            pontos += 5
            motivos.append(f"⏳ Vencimento próximo ({dias_venc} dias restantes)")
            testes.append("Acompanhar planejamento do encerramento da vigência")

    # 6. Monitoramento de Aditamentos (Lei 14.133/21) - Apenas Contratos
    if tipo == "Contratos":
        perc_adit = registro.get("perc_aditivo")
        eh_adit = registro.get("eh_aditivo", False)
        qtd_termos = registro.get("qtd_termos_aditivos", 0)

        if perc_adit is not None:
            if perc_adit > 50.0:
                pontos += 30
                motivos.append(f"⚠️ Aditivo financeiro expressivo (+{perc_adit:.1f}%), acima de 50%")
                testes.append("Auditar termo aditivo: extrapolação do teto legal de 50% para reformas/edifícios")
            elif perc_adit > 25.0:
                pontos += 15
                motivos.append(f"⚠️ Aditivo financeiro relevante (+{perc_adit:.1f}%), acima de 25%")
                testes.append("Auditar termo aditivo: verificar enquadramento legal e justificativa técnica (> 25%)")
            elif perc_adit > 0:
                pontos += 5
                motivos.append(f"Contrato com aditamento financeiro (+{perc_adit:.1f}%)")

        if eh_adit or qtd_termos > 0:
            pontos += 5
            motivos.append(f"Termo Aditivo formal identificado ({qtd_termos} aditivo(s) no PNCP)")
            testes.append("Revisar motivação formal e anexos do termo aditivo")

    # 7. Concentração de Fornecedor
    fornecedor = texto(registro.get("fornecedor"), "").strip()
    if fornecedor and not contexto.empty and "fornecedor" in contexto.columns:
        fornecedores = contexto["fornecedor"].fillna("").astype(str).str.strip().str.casefold()
        qtd = int((fornecedores == fornecedor.casefold()).sum())
        limite_concentracao = max(3, int(len(contexto) * 0.10))

        if qtd >= limite_concentracao:
            pontos += 10
            motivos.append(f"Fornecedor concentrado ({qtd} ocorrências)")
            testes.append("Avaliar concentração e eventual direcionamento")

    pontos = max(0, min(100, int(pontos)))
    if pontos >= 60:
        nivel = "🔴 ALTO"
    elif pontos >= 30:
        nivel = "🟡 MÉDIO"
    else:
        nivel = "🟢 BAIXO"

    if not motivos:
        motivos = ["Sem alertas automáticos."]

    return {
        "pontos": pontos,
        "nivel": nivel,
        "motivos": motivos,
        "testes": testes,
        "outlier": outlier,
    }


# ============================================================
# NLP / SIMILARIDADE SEMÂNTICA OTIMIZADA
# ============================================================

STOPWORDS_CONTRATOS = {
    "a", "o", "os", "as", "um", "uma", "uns", "umas", "de", "do", "da", "dos", "das",
    "em", "no", "na", "nos", "nas", "por", "pelo", "pela", "pelos", "pelas", "com",
    "para", "que", "e", "ou", "ao", "aos", "à", "às", "sob", "sobre", "entre", "se",
    "contrato", "contratacao", "contratada", "contratante", "celebrado", "celebram",
    "termo", "referencia", "especificacoes", "contidas", "anexo", "anexos", "edital",
    "dispensa", "eletronica", "processo", "administrativo", "objeto", "presente",
    "visa", "visando", "atender", "atendimento", "prestacao", "servico", "servicos",
    "empresa", "especializada", "prefeitura", "municipio", "municipal", "secretaria",
    "rio", "pedras", "sp", "conforme", "observadas", "condicoes", "exigencias",
    "estabelecidas", "leis", "lei", "artigo", "art", "vigor", "fornecimento",
    "aquisicao", "locacao", "realizacao", "execucao", "ramo", "me", "epp", "ltda",
    "sa", "cnpj", "cpf", "compromisso", "governo", "hora", "horas", "minutos",
    "dia", "dias", "inicio", "duracao",
}


def texto_nlp(valor: Any) -> str:
    """Normaliza o texto do objeto e remove jargões burocráticos."""
    s = texto(valor, "").lower()

    s = re.sub(r"[áàãâä]", "a", s)
    s = re.sub(r"[éèêë]", "e", s)
    s = re.sub(r"[íìîï]", "i", s)
    s = re.sub(r"[óòõôö]", "o", s)
    s = re.sub(r"[úùûü]", "u", s)
    s = re.sub(r"[ç]", "c", s)

    s = re.sub(r"[^a-z0-9\s]", " ", s)
    palavras = s.split()

    palavras_limpas = [
        p for p in palavras
        if len(p) > 2 and p not in STOPWORDS_CONTRATOS
    ]

    if not palavras_limpas:
        return "objeto_sem_termos_chave"

    return " ".join(palavras_limpas)


@st.cache_data(ttl=CACHE_TTL, show_spinner=False)
def calcular_modelo_tfidf(textos: Tuple[str, ...]):
    """Cria a matriz TF-IDF baseada no núcleo semântico dos objetos."""
    if not SKLEARN_OK or TfidfVectorizer is None or len(textos) < 2:
        return None

    textos_seguros = tuple(texto_nlp(t) for t in textos)
    textos_validos = [t for t in textos_seguros if t != "objeto_sem_termos_chave"]
    if len(set(textos_validos)) < 2:
        return None

    vetor = TfidfVectorizer(
        lowercase=True,
        ngram_range=(1, 2),
        min_df=1,
        max_features=8000,
        sublinear_tf=True,
    )
    try:
        return vetor.fit_transform(textos_seguros)
    except ValueError as exc:
        LOGGER.warning("Não foi possível criar modelo TF-IDF: %s", exc)
        return None


def similares(
    dados_processados: List[Dict[str, Any]],
    idx: int,
    limite: int = 5,
    corte_minimo: float = 0.15,
) -> pd.DataFrame:
    """Retorna registros semanticamente semelhantes filtrando ruídos burocráticos."""
    if not SKLEARN_OK or len(dados_processados) < 2 or idx < 0 or idx >= len(dados_processados):
        return pd.DataFrame()

    objetos = tuple(r.get("objeto") for r in dados_processados)
    matriz = calcular_modelo_tfidf(objetos)
    if matriz is None:
        return pd.DataFrame()

    try:
        scores = (matriz @ matriz[idx].T).toarray().ravel()
    except (IndexError, ValueError) as exc:
        LOGGER.warning("Erro no cálculo de similaridade: %s", exc)
        return pd.DataFrame()

    ordem = scores.argsort()[::-1]
    saida = []

    for j in ordem:
        j = int(j)
        if j == idx:
            continue

        score_val = float(scores[j])
        if score_val < corte_minimo:
            continue

        registro = dados_processados[j]
        saida.append({
            "Similaridade": f"{score_val * 100:.1f}%",
            "Número": texto(registro.get("numero")),
            "Objeto": texto(registro.get("objeto")),
            "Fornecedor": texto(registro.get("fornecedor")),
            "Valor": texto(registro.get("valor")),
        })
        if len(saida) >= limite:
            break

    return pd.DataFrame(saida)


# ============================================================
# EXPORTAÇÃO WORD - DASHBOARD PRINCIPAL (SEM SEÇÃO 3)
# ============================================================

def _word_text(valor: Any, padrao: str = "N/D") -> str:
    """Converte valores para texto adequado ao Word."""
    if valor_vazio(valor):
        return padrao

    if isinstance(valor, float):
        if pd.isna(valor):
            return padrao
        return f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    return texto(valor, padrao)


def _adicionar_tabela_word(
    doc: Document,
    df: pd.DataFrame,
    titulo: Optional[str] = None,
    max_linhas: int = 500,
) -> None:
    """Adiciona um DataFrame ao Word em formato de tabela."""
    if titulo:
        p = doc.add_paragraph()
        run = p.add_run(titulo)
        run.bold = True
        run.font.size = Pt(12)

    if df is None or df.empty:
        doc.add_paragraph("Nenhum dado disponível.")
        return

    tabela_df = df.copy().head(max_linhas)

    tabela = doc.add_table(
        rows=1,
        cols=len(tabela_df.columns),
    )
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.style = "Table Grid"

    for j, coluna in enumerate(tabela_df.columns):
        celula = tabela.rows[0].cells[j]
        celula.text = str(coluna)
        for paragrafo in celula.paragraphs:
            for run in paragrafo.runs:
                run.bold = True
                run.font.size = Pt(8)

    for _, row in tabela_df.iterrows():
        celulas = tabela.add_row().cells
        for j, valor in enumerate(row.tolist()):
            celulas[j].text = _word_text(valor)
            for paragrafo in celulas[j].paragraphs:
                for run in paragrafo.runs:
                    run.font.size = Pt(7)

    if len(df) > max_linhas:
        doc.add_paragraph(
            f"Observação: foram exibidas {max_linhas} linhas de "
            f"{len(df)} registros para manter o documento em tamanho adequado."
        )


def gerar_word_dashboard_principal(
    df_risco: pd.DataFrame,
    tipo: str,
    inicio: Any,
    fim: Any,
    modalidade: str,
    dados_processados: List[Dict[str, Any]],
    riscos: List[Dict[str, Any]],
    valor_total: Any,
) -> bytes:
    """Gera o Word com o resumo executivo: identificação, indicadores, matriz de risco e aviso."""

    if not DOCX_OK:
        raise RuntimeError(
            "A exportação para Word requer a biblioteca "
            "'python-docx'. Adicione 'python-docx' ao requirements.txt "
            "e aguarde o novo deploy."
        )

    doc = Document()

    # Cabeçalho
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("PAINEL DE INTELIGÊNCIA E APOIO AO CONTROLE INTERNO - PNCP")
    run.bold = True
    run.font.size = Pt(16)

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run(PREFEITURA)
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph(
        f"Escopo: {tipo}\n"
        f"Período: {data_br(inicio)} a {data_br(fim)}\n"
        f"Segmentação: {modalidade}\n"
        f"CNPJ do órgão: {CNPJ}\n"
        f"IBGE: {IBGE}"
    )

    # 1. Indicadores Principais
    doc.add_heading("1. Indicadores Principais", level=1)
    altos = sum(r["nivel"] == "🔴 ALTO" for r in riscos)
    medios = sum(r["nivel"] == "🟡 MÉDIO" for r in riscos)

    tabela_indicadores = doc.add_table(rows=4, cols=2)
    tabela_indicadores.style = "Table Grid"
    tabela_indicadores.alignment = WD_TABLE_ALIGNMENT.CENTER

    indicadores = [
        ("Registros Segmentados", str(len(dados_processados))),
        ("🔴 Alto risco", str(altos)),
        ("🟡 Médio risco", str(medios)),
        ("Valor Analisado", moeda(valor_total)),
    ]

    for i, (rotulo, valor) in enumerate(indicadores):
        tabela_indicadores.rows[i].cells[0].text = rotulo
        tabela_indicadores.rows[i].cells[1].text = valor

    # 2. Matriz de Risco Segmentada
    doc.add_heading("2. Matriz de Risco Segmentada", level=1)
    colunas_visualizacao = [
        c for c in (
            "Risco",
            "Score",
            "Número",
            "Tipo",
            "Modalidade",
            "Fornecedor",
            "Valor",
            "Fim Vigência",
            "Aditivo",
            "Objeto",
            "Gatilhos",
        )
        if c in df_risco.columns and (c != "Aditivo" or tipo == "Contratos")
    ]

    _adicionar_tabela_word(
        doc,
        df_risco[colunas_visualizacao] if colunas_visualizacao else df_risco,
        max_linhas=500,
    )

    # Aviso ao Controle Interno
    doc.add_paragraph(
        "Aviso: os alertas e indicadores apresentados neste dashboard "
        "são ferramentas automatizadas de apoio ao Controle Interno. "
        "Os resultados devem ser validados por análise técnica, "
        "jurídica e/ou auditoria humana antes de qualquer conclusão."
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


# ============================================================
# EXPORTAÇÃO EXCEL
# ============================================================

def ajustar_planilha_excel(worksheet: Any) -> None:
    """Formatação básica da planilha."""
    worksheet.freeze_panes = "A2"

    if worksheet.max_row >= 1:
        worksheet.auto_filter.ref = worksheet.dimensions

    for coluna in worksheet.columns:
        if not coluna:
            continue
        letra = coluna[0].column_letter
        maior = 0
        for celula in coluna:
            valor = celula.value
            tamanho = 0 if valor is None else len(str(valor))
            maior = max(maior, tamanho)

        worksheet.column_dimensions[letra].width = min(55, max(12, maior + 2))


def gerar_excel(
    df: pd.DataFrame,
    tipo: str,
    df_risco: pd.DataFrame,
) -> bytes:
    """Gera arquivo Excel para download."""
    buffer = io.BytesIO()

    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df_risco.to_excel(writer, sheet_name="Matriz de Risco", index=False)
        df.to_excel(writer, sheet_name="Dados Brutos", index=False)
        for worksheet in writer.book.worksheets:
            ajustar_planilha_excel(worksheet)

    buffer.seek(0)
    return buffer.getvalue()


# ============================================================
# CONSULTAS ESPECÍFICAS
# ============================================================

def consultar_contratos(inicio: Any, fim: Any) -> Tuple[List[Dict[str, Any]], List[str]]:
    """Consulta contratos."""
    params = {
        "dataInicial": inicio.strftime("%Y%m%d"),
        "dataFinal": fim.strftime("%Y%m%d"),
        "cnpjOrgao": CNPJ,
        "tamanhoPagina": PAGE_CONTRATOS,
    }
    registros, _, _, erros = consultar(f"{BASE_CONSULTA}/contratos", params, MAX_PAGINAS)
    return registros, erros


def consultar_atas(inicio: Any, fim: Any) -> Tuple[List[Dict[str, Any]], List[str]]:
    """Consulta atas de registro de preços."""
    params = {
        "dataInicial": inicio.strftime("%Y%m%d"),
        "dataFinal": fim.strftime("%Y%m%d"),
        "cnpj": CNPJ,
        "tamanhoPagina": PAGE_ATAS,
    }
    registros, _, _, erros = consultar(f"{BASE_CONSULTA}/atas", params, MAX_PAGINAS)
    return registros, erros


def consultar_editais(
    inicio: Any,
    fim: Any,
    progress_callback: Any = None,
) -> Tuple[List[Dict[str, Any]], List[str]]:
    """Consulta editais/contratações por modalidade."""
    params_base = {
        "dataInicial": inicio.strftime("%Y%m%d"),
        "dataFinal": fim.strftime("%Y%m%d"),
        "cnpj": CNPJ,
        "tamanhoPagina": PAGE_EDITAIS,
    }

    todos_registros: List[Dict[str, Any]] = []
    erros: List[str] = []
    total_modalidades = len(MODALIDADES_PNCP)

    for posicao, modalidade in enumerate(MODALIDADES_PNCP, start=1):
        if progress_callback:
            progress_callback(posicao / total_modalidades, modalidade)

        params = dict(params_base)
        params["codigoModalidadeContratacao"] = modalidade

        try:
            registros, _, _, erros_pagina = consultar(
                f"{BASE_CONSULTA}/contratacoes/publicacao",
                params,
                MAX_PAGINAS,
            )
            if registros:
                todos_registros.extend(registros)
            if erros_pagina:
                erros.extend([f"Modalidade {modalidade}: {erro}" for erro in erros_pagina])

        except Exception as exc:
            mensagem = f"Modalidade {modalidade}: {exc}"
            LOGGER.warning(mensagem)
            erros.append(mensagem)

    return (todos_registros, erros)


# ============================================================
# ESTADO DO STREAMLIT
# ============================================================

def dt_date(ano: int, mes: int, dia: int):
    """Helper para criação de date."""
    import datetime as dt
    return dt.date(ano, mes, dia)


def inicializar_estado() -> None:
    """Inicializa session_state."""
    hoje = pd.Timestamp.today().date()
    defaults = {
        "df": None,
        "tipo": TIPOS[0],
        "inicio": dt_date(2026, 1, 1),
        "fim": hoje,
    }

    for chave, valor in defaults.items():
        if chave not in st.session_state:
            st.session_state[chave] = valor


# ============================================================
# VALIDAÇÃO DE PERÍODO
# ============================================================

def validar_periodo(inicio: Any, fim: Any) -> Optional[str]:
    """Valida o período de consulta."""
    if inicio is None or fim is None:
        return "Informe as datas inicial e final."

    if inicio > fim:
        return "A data inicial não pode ser posterior à data final."

    return None


# ============================================================
# INTERFACE
# ============================================================

def main() -> None:

    st.set_page_config(
        page_title="Controle Interno — PNCP Rio das Pedras",
        page_icon="🛡️",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    aplicar_estilo()
    inicializar_estado()

    # Cabeçalho
    st.title("🏛️ Inteligência de Controle Interno PNCP")
    st.info(
        "ℹ️ Sistema de alertas automáticos. "
        "Os resultados são indicadores de apoio e "
        "devem ser validados por análise e auditoria humana."
    )

    # Sidebar
    with st.sidebar:
        st.header("⚙️ Parâmetros")
        tipo_selecionado = st.selectbox(
            "Escopo",
            TIPOS,
            index=TIPOS.index(st.session_state.tipo),
        )

        if tipo_selecionado != st.session_state.tipo:
            st.session_state.tipo = tipo_selecionado
            st.session_state.df = None

        inicio = st.date_input("📅 Data inicial", value=st.session_state.inicio)
        fim = st.date_input("📅 Data final", value=st.session_state.fim)

        st.session_state.inicio = inicio
        st.session_state.fim = fim

        erro_periodo = validar_periodo(inicio, fim)
        if erro_periodo:
            st.error(erro_periodo)

        buscar_btn = st.button(
            "🔎 Carregar dados",
            type="primary",
            use_container_width=True,
            disabled=erro_periodo is not None,
        )

        st.markdown("---")
        st.caption(f"Município: {MUNICIPIO}")
        st.caption(f"CNPJ: {CNPJ}")
        st.caption(f"IBGE: {IBGE}")
        st.caption(f"Máx. páginas por consulta: {MAX_PAGINAS}")

        if not SKLEARN_OK:
            st.warning(
                "Scikit-Learn não instalado. "
                "A análise semântica ficará desativada."
            )

    # Consulta
    if buscar_btn:
        try:
            todas_regs: List[Dict[str, Any]] = []
            erros_consulta: List[str] = []

            with st.spinner("Consultando o PNCP..."):
                tipo_consulta = tipo_selecionado

                if tipo_consulta == "Contratos":
                    registros, erros = consultar_contratos(inicio, fim)
                    todas_regs.extend(registros)
                    erros_consulta.extend(erros)

                elif tipo_consulta == "Atas de Registro de Preços":
                    registros, erros = consultar_atas(inicio, fim)
                    todas_regs.extend(registros)
                    erros_consulta.extend(erros)

                else:
                    barra = st.progress(0, text="🔎 Preparando varredura das modalidades...")

                    def atualizar_progress(progresso: float, modalidade: int) -> None:
                        barra.progress(
                            progresso,
                            text=f"🔎 Varrendo modalidade PNCP cód. {modalidade}...",
                        )

                    registros, erros = consultar_editais(inicio, fim, atualizar_progress)
                    todas_regs.extend(registros)
                    erros_consulta.extend(erros)
                    barra.empty()

            df = normalizar_pncp(todas_regs)
            df = deduplicar(df)

            if tipo_consulta != "Editais e Avisos de Contratações":
                df = filtrar_cnpj(df)

            inicio_ts = pd.Timestamp(inicio).normalize()
            fim_ts = pd.Timestamp(fim).normalize()
            df = filtrar_periodo(df, inicio_ts, fim_ts)

            st.session_state.df = df
            st.session_state.tipo = tipo_consulta

            if erros_consulta:
                st.warning("A consulta terminou com alguns avisos. Verifique os detalhes abaixo.")
                with st.expander("⚠️ Detalhes técnicos da consulta"):
                    for erro in erros_consulta[:30]:
                        st.write(f"- {erro}")
                    if len(erros_consulta) > 30:
                        st.caption(f"{len(erros_consulta) - 30} avisos adicionais foram omitidos.")

            st.rerun()

        except Exception as exc:
            LOGGER.exception("Erro durante consulta PNCP")
            st.error(f"❌ Erro na consulta ao PNCP: {exc}")
            st.stop()

    # Dados
    df_bruto = st.session_state.df
    tipo_atual = st.session_state.tipo

    if df_bruto is None:
        st.info("👈 Defina os parâmetros na barra lateral e clique em **Carregar dados**.")
        st.stop()

    if df_bruto.empty:
        st.warning("Nenhum dado retornado para o filtro aplicado.")
        st.stop()

    # Conversão dos registros
    dados_totais = [dados_registro(row, tipo_atual) for row in df_bruto.to_dict("records")]

    if not dados_totais:
        st.warning("Não foi possível estruturar os registros retornados.")
        st.stop()

    # Segmentação
    st.markdown("---")
    st.subheader("🗂️ Segmentação Estratégica")

    todas_modalidades = sorted({
        d["modalidade"] for d in dados_totais
        if d["modalidade"] and d["modalidade"] != "N/D"
    })

    opcao_geral = "Visão Geral (Todas as Modalidades)"

    if todas_modalidades:
        mod_escolhida = st.selectbox(
            "Filtrar e Segmentar por Modalidade:",
            [opcao_geral, *todas_modalidades],
        )
    else:
        mod_escolhida = opcao_geral

    if mod_escolhida != opcao_geral:
        indices_ativos = [
            i for i, registro in enumerate(dados_totais)
            if registro["modalidade"] == mod_escolhida
        ]
    else:
        indices_ativos = list(range(len(dados_totais)))

    if not indices_ativos:
        st.warning(f"Sem registros para a modalidade: {mod_escolhida}")
        st.stop()

    df_filtrado = df_bruto.iloc[indices_ativos].reset_index(drop=True)
    dados_processados = [dados_totais[i] for i in indices_ativos]
    contexto = pd.DataFrame(dados_processados)

    riscos = [
        calcular_risco(registro, contexto, tipo_atual)
        for registro in dados_processados
    ]

    # Indicadores Gerais
    altos = sum(r["nivel"] == "🔴 ALTO" for r in riscos)
    medios = sum(r["nivel"] == "🟡 MÉDIO" for r in riscos)
    baixos = sum(r["nivel"] == "🟢 BAIXO" for r in riscos)

    if "valor_num" in contexto.columns and not contexto["valor_num"].dropna().empty:
        valor_total = contexto["valor_num"].sum()
    else:
        valor_total = None

    k1, k2, k3, k4 = st.columns(4)
    k1.metric("Registros Segmentados", len(dados_processados))
    k2.metric("🔴 Alto risco", altos)
    k3.metric("🟡 Médio risco", medios)
    k4.metric("Valor Analisado", moeda(valor_total))

    # Matriz de Risco Estruturada
    rows = []
    for i_local, risco in enumerate(riscos):
        registro = dados_processados[i_local]
        
        fim_vig = registro["data_fim_vigencia"]
        dias_v = registro["dias_para_vencer"]
        if dias_v is not None:
            if dias_v < 0:
                vig_status = f"{fim_vig} (Vencido há {abs(dias_v)}d)"
            else:
                vig_status = f"{fim_vig} ({dias_v}d restantes)"
        else:
            vig_status = fim_vig

        perc_a = registro["perc_aditivo"]
        qtd_t = registro.get("qtd_termos_aditivos", 0)
        
        if perc_a is not None:
            adit_status = f"+{perc_a:.1f}%"
        elif qtd_t > 0:
            adit_status = f"{qtd_t} termo(s) no PNCP"
        elif registro["eh_aditivo"]:
            adit_status = "Aditivo formal"
        else:
            adit_status = "Inicial (0%)"

        rows.append({
            "Índice": i_local,
            "Risco": risco["nivel"],
            "Score": risco["pontos"],
            "Número": registro["numero"],
            "Tipo": registro["tipo_instrumento"],
            "Modalidade": registro["modalidade"],
            "Fornecedor": registro["fornecedor"],
            "Valor": registro["valor"],
            "Fim Vigência": vig_status,
            "Aditivo": adit_status,
            "Objeto": registro["objeto"],
            "Gatilhos": "; ".join(risco["motivos"]),
        })

    df_risco = pd.DataFrame(rows)
    if not df_risco.empty:
        df_risco = df_risco.sort_values(
            ["Score", "Índice"],
            ascending=[False, True],
        ).reset_index(drop=True)

    # Abas da Aplicação
    st.markdown("---")
    tab1, tab_vigencia, tab2 = st.tabs([
        "🚦 Matriz de Risco Segmentada",
        "⏳ Vigência e Aditivos Legais",
        "📋 Auditoria e Semântica Individual",
    ])

    # ========================================================
    # ABA 1 - MATRIZ DE RISCO PRINCIPAL (COM OPÇÕES DE DOWNLOAD)
    # ========================================================
    with tab1:
        colunas_visualizacao = [
            c for c in (
                "Risco",
                "Score",
                "Número",
                "Tipo",
                "Modalidade",
                "Fornecedor",
                "Valor",
                "Fim Vigência",
                "Aditivo",
                "Objeto",
                "Gatilhos",
            )
            if c in df_risco.columns and (c != "Aditivo" or tipo_atual == "Contratos")
        ]
        
        st.info(
            "💡 **Entenda os Alertas de Risco:**\n\n"
            "**• Valor anômalo acima do limite IQR:** O sistema utiliza o método estatístico de Intervalo "
            "Interquartil (IQR) para identificar discrepâncias financeiras. Um alerta é gerado quando o "
            "valor foge significativamente do padrão dos demais itens do mesmo segmento (outlier). Indica a "
            "necessidade de revisar a formação de preços e pesquisa de mercado.\n\n"
            "**• Fornecedor concentrado:** Ocorre quando uma mesma empresa acumula um volume alto de "
            "contratações (acima de 10% do total analisado). O alerta sugere avaliar se há dependência "
            "econômica do fornecedor, riscos à continuidade do fornecimento ou indícios de direcionamento."
        )

        st.dataframe(
            df_risco[colunas_visualizacao],
            use_container_width=True,
            hide_index=True,
        )

        try:
            excel = gerar_excel(
                df_filtrado,
                tipo_atual,
                df_risco,
            )
            nome_arquivo = f"Auditoria_{slug(mod_escolhida)}.xlsx"

            col_download1, col_download2 = st.columns(2)

            with col_download1:
                st.download_button(
                    "📊 Baixar Matriz Deste Segmento (Excel)",
                    data=excel,
                    file_name=nome_arquivo,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )

            if DOCX_OK:
                try:
                    word_aba1 = gerar_word_dashboard_principal(
                        df_risco=df_risco,
                        tipo=tipo_atual,
                        inicio=inicio,
                        fim=fim,
                        modalidade=mod_escolhida,
                        dados_processados=dados_processados,
                        riscos=riscos,
                        valor_total=valor_total,
                    )
                    nome_word_aba1 = f"Dashboard_Matriz_Risco_{slug(mod_escolhida)}.docx"

                    with col_download2:
                        st.download_button(
                            "📝 Baixar Dashboard em Word",
                            data=word_aba1,
                            file_name=nome_word_aba1,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )
                except Exception as exc:
                    LOGGER.exception("Erro ao gerar Word do dashboard principal")
                    with col_download2:
                        st.error(f"Não foi possível gerar o Word: {exc}")
            else:
                with col_download2:
                    st.warning("python-docx não instalado para exportação Word.")

        except Exception as exc:
            LOGGER.exception("Erro ao gerar Excel")
            st.error(f"Não foi possível gerar o Excel: {exc}")

    # ========================================================
    # ABA 2 - VIGÊNCIA E ADITIVOS LEGAIS
    # ========================================================
    with tab_vigencia:
        if tipo_atual == "Contratos":
            st.subheader("⏱️ Monitoramento Preditivo de Vigência & Limites da Lei 14.133/21")
            st.markdown(
                "Rastreamento de contratos com prazo expirando para planejamento de novas contratações "
                "e verificação de aditamentos formais e financeiros (+25% em geral / +50% reformas)."
            )
        else:
            st.subheader("⏱️ Monitoramento Preditivo de Vigência")
            st.markdown(
                "Rastreamento de prazos expirando para planejamento tempestivo de novas licitações."
            )

        vencendo_30 = [r for r in dados_processados if r.get("dias_para_vencer") is not None and 0 <= r["dias_para_vencer"] <= 30]
        vencendo_90 = [r for r in dados_processados if r.get("dias_para_vencer") is not None and 30 < r["dias_para_vencer"] <= 90]
        
        # Preparação comum da tabela de vigências
        tabela_vig = []
        for r in dados_processados:
            dias_rest = r.get("dias_para_vencer")
            if dias_rest is None:
                status_prazo = "⚪ Não informado"
            elif dias_rest < 0:
                status_prazo = f"⚫ Encerrado ({abs(dias_rest)}d atrás)"
            elif dias_rest <= 30:
                status_prazo = f"🔴 Crítico ({dias_rest} dias)"
            elif dias_rest <= 90:
                status_prazo = f"🟡 Atenção ({dias_rest} dias)"
            else:
                status_prazo = f"🟢 Regular ({dias_rest} dias)"

            tabela_vig.append({
                "Número": r["numero"],
                "Tipo": r["tipo_instrumento"],
                "Fornecedor": r["fornecedor"],
                "Data Assinatura": r["data"],
                "Término Vigência": r["data_fim_vigencia"],
                "Status do Prazo": status_prazo,
                "Valor": r["valor"],
                "Objeto": r["objeto"],
            })
        df_vig_show = pd.DataFrame(tabela_vig)

        if tipo_atual == "Contratos":
            total_aditivos = [
                r for r in dados_processados 
                if r.get("eh_aditivo", False) 
                or (r.get("qtd_termos_aditivos", 0) > 0)
                or (r.get("perc_aditivo") is not None and r["perc_aditivo"] > 0)
            ]

            m1, m2, m3 = st.columns(3)
            m1.metric("🚨 Vencem em até 30 dias", len(vencendo_30))
            m2.metric("⚠️ Vencem entre 31 e 90 dias", len(vencendo_90))
            m3.metric("📈 Termos Aditivos Identificados", len(total_aditivos))

            # ========================================================
            # DIAGNÓSTICO REAL DA CONSULTA DE ADITIVOS
            # ========================================================
            diagnosticos = [
                r.get("diagnostico_aditivo", {})
                for r in dados_processados
            ]

            total_contratos = len(diagnosticos)
            qtd_encontrados = sum(
                d.get("status") == "REGISTROS_ENCONTRADOS"
                for d in diagnosticos
            )
            qtd_sem_registros = sum(
                d.get("status") == "SEM_REGISTROS"
                for d in diagnosticos
            )
            qtd_sem_identificadores = sum(
                d.get("status") == "SEM_IDENTIFICADORES"
                for d in diagnosticos
            )
            qtd_404 = sum(
                d.get("status") == "HTTP_404"
                for d in diagnosticos
            )
            qtd_erros = sum(
                d.get("status") in {"ERRO_CONSULTA", "RESPOSTA_INVALIDA"}
                for d in diagnosticos
            )

            st.markdown("### 🔎 Diagnóstico da Consulta de Aditivos")

            d1, d2, d3, d4, d5 = st.columns(5)
            d1.metric("Contratos analisados", total_contratos)
            d2.metric("🟢 Com termos", qtd_encontrados)
            d3.metric("⚪ Sem termos", qtd_sem_registros)
            d4.metric("🟡 Sem identificadores", qtd_sem_identificadores)
            d5.metric("🔴 Falhas", qtd_erros + qtd_404)

            if qtd_encontrados > 0:
                st.success(
                    f"✅ O PNCP retornou termos para {qtd_encontrados} contrato(s). "
                    "Esses registros foram considerados na análise de aditamentos."
                )

            if qtd_sem_registros > 0:
                st.info(
                    f"ℹ️ {qtd_sem_registros} contrato(s) foram consultados com sucesso "
                    "e não apresentaram termos retornados pelo PNCP."
                )
                
                tabela_sem_termos = []
                for r in dados_processados:
                    if r.get("diagnostico_aditivo", {}).get("status") == "SEM_REGISTROS":
                        tabela_sem_termos.append({
                            "Contrato": r.get("numero", "N/D"),
                            "Ano": r.get("ano_contrato", "N/D"),
                            "Sequencial": r.get("sequencial_contrato", "N/D"),
                            "Resultado": "Nenhum termo no PNCP",
                        })
                        
                st.dataframe(
                    pd.DataFrame(tabela_sem_termos),
                    use_container_width=True,
                    hide_index=True,
                )
                
                st.markdown(
                    "ℹ️ **A ausência de termo no PNCP não comprova, isoladamente, a inexistência de "
                    "aditamento administrativo. Recomenda-se confrontar o resultado com os processos, "
                    "contratos e documentos mantidos pela Prefeitura.**"
                )

            if qtd_sem_identificadores > 0:
                st.warning(
                    f"⚠️ {qtd_sem_identificadores} contrato(s) não puderam ter os termos "
                    "consultados porque anoContrato e/ou sequencialContrato não foram "
                    "encontrados no registro retornado pelo PNCP."
                )

            if qtd_404 > 0:
                st.warning(
                    f"🟠 {qtd_404} consulta(s) retornaram HTTP 404. Isso significa que o "
                    "recurso de termos não foi localizado pelo endpoint consultado; não "
                    "é correto concluir apenas por isso que não existem aditivos."
                )

            if qtd_erros > 0:
                st.error(
                    f"🔴 {qtd_erros} consulta(s) apresentaram erro ou resposta inválida. "
                    "Não é possível usar esses casos para afirmar ausência de aditivos."
                )

            with st.expander("📋 Detalhes técnicos das consultas", expanded=False):
                resumo_diagnostico = pd.DataFrame([
                    {"Situação": "Termos encontrados", "Quantidade": qtd_encontrados},
                    {"Situação": "Consulta realizada sem termos", "Quantidade": qtd_sem_registros},
                    {"Situação": "Identificadores ausentes", "Quantidade": qtd_sem_identificadores},
                    {"Situação": "HTTP 404", "Quantidade": qtd_404},
                    {"Situação": "Erro/resposta inválida", "Quantidade": qtd_erros},
                ])
                st.dataframe(
                    resumo_diagnostico,
                    use_container_width=True,
                    hide_index=True,
                )

                detalhes_diagnostico = []
                for r in dados_processados:
                    diag = r.get("diagnostico_aditivo", {})
                    status = diag.get("status", "NAO_CONSULTADO")
                    rotulos = {
                        "REGISTROS_ENCONTRADOS": "🟢 Termo(s) encontrado(s)",
                        "SEM_REGISTROS": "⚪ Consulta realizada — nenhum termo",
                        "SEM_IDENTIFICADORES": "🟡 Identificadores insuficientes",
                        "HTTP_404": "🟠 HTTP 404",
                        "ERRO_CONSULTA": "🔴 Erro na consulta",
                        "RESPOSTA_INVALIDA": "🔴 Resposta inválida",
                        "NAO_CONSULTADO": "⚪ Não consultado",
                    }
                    detalhes_diagnostico.append({
                        "Contrato": r.get("numero", "N/D"),
                        "Ano": r.get("ano_contrato", "N/D"),
                        "Sequencial": r.get("sequencial_contrato", "N/D"),
                        "Situação": rotulos.get(status, status),
                        "HTTP": diag.get("http_status", ""),
                        "Termos": len(diag.get("termos", [])),
                        "Mensagem": diag.get("mensagem", ""),
                    })

                st.dataframe(
                    pd.DataFrame(detalhes_diagnostico),
                    use_container_width=True,
                    hide_index=True,
                )

            st.markdown("---")

            sub_tab_vig, sub_tab_adit = st.tabs(["📅 Linha do Tempo de Vigências", "📊 Auditoria de Aditamentos"])

            with sub_tab_vig:
                st.dataframe(df_vig_show, use_container_width=True, hide_index=True)

            with sub_tab_adit:
                tabela_adit = []
                for r in dados_processados:
                    perc_a = r.get("perc_aditivo")
                    eh_adit = r.get("eh_aditivo", False)
                    qtd_t = r.get("qtd_termos_aditivos", 0)
                    v_ini = r.get("valor_inicial_num")
                    v_fim = r.get("valor_num")

                    if eh_adit or qtd_t > 0 or (perc_a is not None and perc_a > 0):
                        if perc_a is not None and perc_a > 50:
                            status_legal = "🔴 > 50% (Risco Alto)"
                        elif perc_a is not None and perc_a > 25:
                            status_legal = "🟡 > 25% (Risco Médio)"
                        elif perc_a is not None and perc_a > 0:
                            status_legal = "🟢 Aditivo Finance Regular"
                        elif qtd_t > 0:
                            status_legal = f"ℹ️ {qtd_t} Termo(s) Aditivo(s) no PNCP"
                        else:
                            status_legal = "ℹ️ Aditivo de Prazo / Qualitativo"

                        tabela_adit.append({
                            "Número": r["numero"],
                            "Tipo de Instrumento": r["tipo_instrumento"],
                            "Termos no PNCP": qtd_t,
                            "Fornecedor": r["fornecedor"],
                            "Valor Inicial": moeda(v_ini) if v_ini else "N/D",
                            "Valor Atualizado": moeda(v_fim),
                            "Variação (%)": f"+{perc_a:.2f}%" if perc_a is not None else "0.00%",
                            "Avaliação de Conformidade": status_legal,
                            "Objeto": r["objeto"],
                        })

                if tabela_adit:
                    df_adit_show = pd.DataFrame(tabela_adit)
                    st.dataframe(
                        df_adit_show,
                        use_container_width=True,
                        hide_index=True,
                    )
                else:
                    if (
                        qtd_sem_registros > 0
                        and qtd_sem_identificadores == 0
                        and qtd_404 == 0
                        and qtd_erros == 0
                    ):
                        st.success(
                            "✅ As consultas aos termos foram realizadas com sucesso. "
                            "Nenhum termo aditivo foi retornado pelo PNCP para os "
                            "contratos analisados."
                        )
                    else:
                        st.warning(
                            "⚠️ Nenhum aditivo foi exibido na tabela, mas isso não "
                            "permite concluir que não existam aditivos. Consulte o "
                            "diagnóstico acima para verificar as consultas que não "
                            "puderam ser concluídas."
                        )

        else:
            # Caso não seja Contratos (Atas ou Editais), renderiza vigência diretamente sem a lógica de aditivos
            m1, m2 = st.columns(2)
            m1.metric("🚨 Vencem em até 30 dias", len(vencendo_30))
            m2.metric("⚠️ Vencem entre 31 e 90 dias", len(vencendo_90))

            st.markdown("---")
            st.markdown("### 📅 Linha do Tempo de Vigências")
            st.dataframe(df_vig_show, use_container_width=True, hide_index=True)


    # ========================================================
    # ABA 3 - AUDITORIA E SEMÂNTICA INDIVIDUAL
    # ========================================================
    with tab2:
        f1, f2 = st.columns(2)

        filtro_risco = f1.multiselect(
            "Nível de Risco Alvo",
            ["🔴 ALTO", "🟡 MÉDIO", "🟢 BAIXO"],
            default=["🔴 ALTO", "🟡 MÉDIO"],
        )

        mostrar_outlier = f2.checkbox("Apenas anomalias financeiras")

        idx_auditoria = [
            i for i, risco in enumerate(riscos)
            if risco["nivel"] in filtro_risco and (not mostrar_outlier or risco["outlier"])
        ]

        if not idx_auditoria:
            st.warning("Ajuste os filtros de priorização acima para investigar um dossiê.")
        else:
            mapa = {}
            for i in idx_auditoria:
                risco = riscos[i]
                registro = dados_processados[i]
                chave = f"[{risco['nivel']} | Score {risco['pontos']}] {registro['numero']} — {registro['valor']}"

                if chave in mapa:
                    contador = 2
                    chave_base = chave
                    while chave in mapa:
                        chave = f"{chave_base} (#{contador})"
                        contador += 1

                mapa[chave] = i

            escolhido = st.selectbox(
                "Selecione o dossiê da contratação:",
                list(mapa.keys()),
            )

            i_alvo = mapa[escolhido]
            registro = dados_processados[i_alvo]
            risco = riscos[i_alvo]

            with st.container():
                col_a, col_b, col_c = st.columns(3)
                col_a.metric("Score de Risco", f"{risco['pontos']} pts")
                col_b.metric("Volume", registro["valor"])
                col_c.metric("Enquadramento Legal", registro["modalidade"])

                st.info(f"**Objeto Descrito:** {registro['objeto']}")

            with st.expander("🔎 Justificativa Analítica e Testes", expanded=True):
                st.markdown("**Variáveis de Acionamento:**")
                for motivo in risco["motivos"]:
                    st.write(f"- {motivo}")

                if risco["testes"]:
                    st.markdown("**Roteiro de Auditoria Sugerido:**")
                    for teste in risco["testes"]:
                        st.write(f"- {teste}")
                else:
                    st.caption("Nenhum teste específico foi sugerido automaticamente.")

            with st.expander("📄 Informações do Registro", expanded=False):
                info_col1, info_col2 = st.columns(2)
                info_col1.write(f"**Controle PNCP:** {registro['controle']}")
                info_col1.write(f"**Número:** {registro['numero']}")
                info_col1.write(f"**Tipo de Instrumento:** {registro['tipo_instrumento']}")
                info_col1.write(f"**Processo:** {registro['processo']}")
                info_col1.write(f"**Data Assinatura:** {registro['data']}")
                info_col1.write(f"**Fim da Vigência:** {registro['data_fim_vigencia']}")

                info_col2.write(f"**Fornecedor:** {registro['fornecedor']}")
                info_col2.write(f"**CNPJ Fornecedor:** {registro['cnpj_fornecedor'] or 'N/D'}")
                info_col2.write(f"**Situação:** {registro['situacao']}")
                info_col2.write(f"**Modalidade:** {registro['modalidade']}")

            st.markdown("### 🤖 Clusterização Semântica (Apenas neste Segmento)")

            if not SKLEARN_OK:
                st.warning(
                    "Scikit-Learn não está instalado. "
                    "Instale a dependência para habilitar a análise semântica."
                )
            else:
                sim = similares(dados_processados, i_alvo, limite=5, corte_minimo=0.15)
                if sim.empty:
                    st.info("Não foi possível identificar vizinhança semântica relevante para este item.")
                else:
                    st.dataframe(sim, use_container_width=True, hide_index=True)

    # Rodapé
    st.markdown("---")
    st.caption(
        f"Painel de apoio ao Controle Interno — {PREFEITURA}. "
        "Dados públicos consultados no PNCP. "
        "Alertas automatizados não substituem análise técnica, jurídica ou auditoria humana."
    )


# ============================================================
# EXECUÇÃO
# ============================================================

if __name__ == "__main__":
    main()
